# -*- coding: utf-8 -*-
"""
Script tạo 3 file .docx chuyên nghiệp cho đồ án CSDL Phân Tán.
Chạy: python generate_docx.py
"""
import sys, io
if sys.platform == "win32":
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')
    except Exception:
        pass

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def setup_styles(doc):
    """Cau hinh style chuan cho document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    # Set font cho tat ca Heading
    for i in range(1, 5):
        h_name = f'Heading {i}'
        if h_name in doc.styles:
            hs = doc.styles[h_name]
            hs.font.name = 'Times New Roman'
            hs.font.color.rgb = RGBColor(0, 51, 102)
            hs.paragraph_format.space_before = Pt(12)
            hs.paragraph_format.space_after = Pt(6)
            if i == 1:
                hs.font.size = Pt(18)
                hs.font.bold = True
            elif i == 2:
                hs.font.size = Pt(15)
                hs.font.bold = True
            elif i == 3:
                hs.font.size = Pt(13)
                hs.font.bold = True
                hs.font.italic = False

def set_cell_shading(cell, color_hex):
    """Dat mau nen cho cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table):
    """Dat vien cho toan bo table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="4472C4"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="4472C4"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="4472C4"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="4472C4"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="4472C4"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="4472C4"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Tao bang dep voi header co mau nen xanh dam."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '003366')
    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if i % 2 == 1:
                set_cell_shading(cell, 'E8F0FE')
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table

def add_para(doc, text, bold=False, italic=False, size=13, align=None, space_after=6, indent=0):
    """Them mot doan van ban."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def add_bullet(doc, text, level=0, bold_prefix=""):
    """Them dau gach dau dong."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(12)
        run_b.font.name = 'Times New Roman'
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * level)
    return p

def add_title_page(doc, title_lines, subtitle=""):
    """Tao trang bia."""
    for _ in range(4):
        doc.add_paragraph()
    for line in title_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(80, 80, 80)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3):
        doc.add_paragraph()

def add_horizontal_line(doc):
    """Them duong ke ngang."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="003366"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


# ============================================================
# FILE 1: PROJECT PROPOSAL
# ============================================================

def create_project_proposal():
    doc = Document()
    setup_styles(doc)

    # --- TRANG BIA ---
    add_title_page(doc,
        ["ĐỀ CƯƠNG DỰ ÁN", "CƠ SỞ DỮ LIỆU PHÂN TÁN"],
        "DISTRIBUTED DATABASE PROJECT PROPOSAL"
    )
    add_para(doc, "Đề tài #15: Bloom Filter Join Optimizer", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Category 3: Tối ưu hóa Truy vấn Phân tán", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "Due Date: [Điền ngày nộp - Tuần 3]", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # --- 1. PROJECT IDENTITY ---
    doc.add_heading('1. Project Identity (Thông tin nhóm)', level=1)
    add_horizontal_line(doc)

    add_styled_table(doc,
        ["Thông tin", "Chi tiết"],
        [
            ["Team Name", "[Điền tên nhóm — Ví dụ: BloomBand / FilterFlow]"],
            ["Thành viên 1", "[Họ và tên — MSSV]"],
            ["Thành viên 2", "[Họ và tên — MSSV]"],
        ],
        col_widths=[4, 12]
    )
    doc.add_paragraph()
    add_para(doc, "Project Title:", bold=True, size=13)
    add_para(doc,
        "Thiết kế và Đánh giá Bộ Tối ưu hóa Phép Nối Phân Tán sử dụng Bloom Filter Semi-Join "
        "cho Hệ thống Subscribers & WebLogs",
        italic=True, size=12, indent=0.5
    )
    add_para(doc,
        "(Bloom Filter Semi-Join Query Optimizer for Distributed Subscriber & WebLog Systems)",
        italic=True, size=11, indent=0.5
    )

    # --- 2. OBJECTIVE ---
    doc.add_heading('2. Objective & Problem Statement (Mục tiêu & Bài toán)', level=1)
    add_horizontal_line(doc)

    add_para(doc, 'The "Why" (Lý do thực hiện):', bold=True, size=13)
    add_para(doc,
        'Trong hệ cơ sở dữ liệu phân tán, việc thực hiện phép nối (JOIN) hai bảng lớn nằm ở '
        'hai Site khác nhau qua đường truyền mạng là một "nút thắt cổ chai" (bottleneck) lớn '
        'về mặt băng thông và thời gian phản hồi.', size=12
    )
    add_para(doc,
        'Nếu truyền toàn bộ dữ liệu thô chưa qua lọc từ Site B sang Site A (phương pháp Naive Join), '
        'lượng băng thông tiêu thụ sẽ rất lớn, trong đó có tới 80% là dữ liệu thừa của những '
        'người dùng không tồn tại ở Site A (guest/anonymous logs).', size=12
    )
    add_para(doc,
        'Dự án này giải quyết bài toán tối ưu hóa băng thông bằng cách áp dụng thuật toán '
        'Bloom Filter Semi-Join để lọc dữ liệu thừa ngay tại Site nguồn trước khi truyền tải '
        'qua mạng. Chúng tôi sẽ phân tích mối quan hệ đánh đổi (trade-off) giữa kích thước '
        'của vector bit (m bits), số lượng hàm băm (k), tỷ lệ lỗi dương tính giả (FPR) '
        'và băng thông mạng tiết kiệm được.', size=12
    )

    add_para(doc, 'Core Logic (Thuật toán cốt lõi) — 5 bước chính:', bold=True, size=13)
    steps = [
        ("Bước 1 — Tạo Bloom Filter tại Site A: ", "Trích xuất toàn bộ khóa nối user_id từ bảng Subscribers (1 triệu dòng), băm qua k hàm băm và ánh xạ vào vector bit kích thước m bits."),
        ("Bước 2 — Truyền Bloom Filter: ", "Gửi vector bit siêu nhỏ này từ Site A sang Site B qua mạng."),
        ("Bước 3 — Lọc dữ liệu tại Site B: ", "Duyệt qua bảng WebLogs (10 triệu dòng) tại Site B. Kiểm tra user_id của từng dòng log với Bloom Filter. Chỉ giữ lại những dòng log có user_id cho kết quả khớp."),
        ("Bước 4 — Truyền kết quả lọc: ", "Site B gửi các dòng log đã lọc (chỉ khoảng 20% dung lượng gốc) ngược lại Site A."),
        ("Bước 5 — Khớp nối cuối cùng: ", "Thực hiện phép Inner Join thực tế giữa bảng Subscribers và các dòng log nhận được để loại bỏ hoàn toàn False Positives, đảm bảo kết quả chính xác 100%."),
    ]
    for prefix, desc in steps:
        add_bullet(doc, desc, bold_prefix=prefix)

    # --- 3. DATASET ---
    doc.add_heading('3. Dataset Specification (Đặc tả dữ liệu)', level=1)
    add_horizontal_line(doc)

    add_para(doc, 'Source: Dữ liệu giả lập có cấu trúc thực tế, sinh tự động bằng Python script '
             '(2_data_generator.py) để kiểm soát tỷ lệ trùng khớp khóa nối giữa hai Site.', size=12)

    add_styled_table(doc,
        ["Site", "Bảng", "Quy mô"],
        [
            ["Site A (Trụ sở)", "Subscribers — Thuê bao trả phí", "1,000,000 rows (~80 MB)"],
            ["Site B (Web Server)", "WebLogs — Nhật ký truy cập", "10,000,000 rows (~1.2 GB)"],
        ],
        col_widths=[4, 6, 5]
    )
    doc.add_paragraph()

    add_para(doc, 'Schema — Subscribers (Site A):', bold=True, size=12)
    add_styled_table(doc,
        ["Cột", "Kiểu dữ liệu", "Mô tả"],
        [
            ["user_id", "PK — String", "Mã định danh duy nhất"],
            ["full_name", "String", "Tên đầy đủ (tiếng Việt)"],
            ["email", "String", "Địa chỉ email"],
            ["plan", "String", "Gói dịch vụ (Free/Bronze/Silver/Gold)"],
            ["monthly_fee", "Float", "Phí thuê bao hàng tháng"],
        ],
        col_widths=[3.5, 3.5, 8]
    )
    doc.add_paragraph()

    add_para(doc, 'Schema — WebLogs (Site B):', bold=True, size=12)
    add_styled_table(doc,
        ["Cột", "Kiểu dữ liệu", "Mô tả"],
        [
            ["log_id", "PK — String", "Mã định danh dòng log"],
            ["user_id", "FK — String", "Mã người dùng thực hiện hành động"],
            ["page", "String", "Trang web truy cập (/home, /dashboard…)"],
            ["method", "String", "Phương thức HTTP (GET, POST)"],
            ["status", "Integer", "Mã trạng thái HTTP (200, 404, 500)"],
            ["timestamp", "DateTime", "Thời gian ghi nhận log"],
        ],
        col_widths=[3.5, 3.5, 8]
    )
    doc.add_paragraph()

    add_para(doc, 'Fragmentation Strategy:', bold=True, size=12)
    add_para(doc,
        'Phân tán ngang (Horizontal Fragmentation) theo vị trí địa lý vật lý: '
        'Bảng Subscribers lưu tại Site A (Trụ sở chính), '
        'bảng WebLogs lưu tại Site B (Web Server tại Data Center).', size=12
    )

    # --- 4. SYSTEM ARCHITECTURE ---
    doc.add_heading('4. System Architecture (Kiến trúc hệ thống)', level=1)
    add_horizontal_line(doc)

    add_para(doc, 'Nodes — Giả lập 2 sites chính:', bold=True, size=12)
    add_bullet(doc, "Nơi khởi tạo truy vấn và nhận kết quả cuối cùng.", bold_prefix="Site A (Client/Trụ sở): ")
    add_bullet(doc, "Nơi lưu giữ lượng nhật ký WebLogs khổng lồ.", bold_prefix="Site B (Server/Chi nhánh): ")

    add_para(doc, 'Communication Layer:', bold=True, size=12)
    add_bullet(doc, "Giao tiếp giữa các tiến trình thông qua Web API (HTTP/REST sử dụng Flask).")
    add_bullet(doc, "Có bộ giám sát lưu lượng mạng (Network Traffic Monitor) để tính toán dung lượng truyền tải thực tế.")

    add_para(doc, 'Storage:', bold=True, size=12)
    add_bullet(doc, "Dữ liệu lưu trữ dạng tập tin CSV cục bộ tại mỗi Site để mô phỏng tính độc lập.")
    add_bullet(doc, "Tải vào bộ nhớ dưới dạng Pandas DataFrame để mô phỏng truy vấn.")

    # --- 5. TECH STACK ---
    doc.add_heading('5. Tech Stack & Implementation Plan (Công nghệ & Kế hoạch)', level=1)
    add_horizontal_line(doc)

    add_styled_table(doc,
        ["Thành phần", "Công nghệ", "Vai trò"],
        [
            ["Ngôn ngữ", "Python 3.8+", "Xử lý mảng và cấu trúc dữ liệu lớn"],
            ["Hash function", "mmh3 (MurmurHash3)", "Hàm băm phi mã hóa, phân bố đều bit"],
            ["Bit storage", "bitarray", "Mảng bit tối ưu bộ nhớ RAM"],
            ["Data processing", "Pandas", "Xử lý, nối và truy vấn dữ liệu lớn"],
            ["Web framework", "Flask", "RESTful API mô phỏng giao tiếp mạng"],
            ["Visualization", "Chart.js + Matplotlib", "Biểu đồ động trên Web Dashboard"],
            ["Deployment", "Flask local server", "Web Dashboard trên trình duyệt"],
        ],
        col_widths=[3.5, 4, 8]
    )

    # --- 6. SUCCESS METRICS ---
    doc.add_heading('6. Success Metrics & Analysis (Chỉ số đo lường)', level=1)
    add_horizontal_line(doc)

    add_para(doc, 'Chỉ số định lượng:', bold=True, size=12)
    add_bullet(doc, "Tính bằng: Bytes_Saved = Bandwidth_Naive − Bandwidth_BF_Semi_Join.", bold_prefix="Băng thông tiết kiệm (Bytes saved): ")
    add_bullet(doc, "Mục tiêu đạt trên 70%.", bold_prefix="Tỷ lệ tiết kiệm (Bandwidth Savings %): ")
    add_bullet(doc, "Tỷ lệ byte mạng tiết kiệm / kích thước BF.", bold_prefix="Savings Leverage: ")
    add_bullet(doc, "So sánh FPR lý thuyết vs thực tế.", bold_prefix="False Positive Rate thực nghiệm: ")

    add_para(doc, 'Kịch bản lỗi (Failure Scenario):', bold=True, size=12)
    add_bullet(doc, "Khi giảm kích thước m quá nhỏ, FPR tăng đột ngột (>30%), Site B gửi nhầm hàng triệu dòng rác. "
               "Phân tích điểm gãy mạng (Break-even Point).", bold_prefix="1. Bộ lọc quá nhỏ: ")
    add_bullet(doc, "Node B bị sập giữa chừng. Site A bắt lỗi Connection Timeout (HTTP 504), "
               "tự động Abort Transaction và rollback trạng thái.", bold_prefix="2. Sập mạng / Kill Node B: ")

    # --- 7. MILESTONES ---
    doc.add_heading('7. Project Milestones (Mốc thời gian)', level=1)
    add_horizontal_line(doc)

    add_styled_table(doc,
        ["Mốc", "Tuần", "Nội dung"],
        [
            ["Milestone 1", "Week 5", "Thiết lập môi trường. Hoàn thành script sinh dữ liệu 1M/10M."],
            ["Milestone 2", "Week 8", "Cài đặt Bloom Filter (Double Hashing + MurmurHash3). Pipeline Semi-Join hoàn chỉnh."],
            ["Milestone 3", "Week 12", "Web Dashboard trực quan. Benchmark các kích thước BF. Hoàn thiện báo cáo."],
        ],
        col_widths=[3, 2, 10]
    )

    # --- LUU FILE ---
    filepath = os.path.join(OUTPUT_DIR, "Project_Proposal_v2.docx")
    doc.save(filepath)
    print(f"  [OK] Da tao: {filepath}")


# ============================================================
# FILE 2: DESIGN DOCUMENT
# ============================================================

def create_design_document():
    doc = Document()
    setup_styles(doc)

    # --- HEADER (không trang bìa riêng — tiết kiệm giấy) ---
    add_para(doc, "DESIGN DOCUMENT", bold=True, size=18,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, 'Bloom Filter Semi-Join Optimizer — "Subscribers & Logs"',
             italic=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Đề tài #15 | Môn: Cơ Sở Dữ Liệu Phân Tán",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_horizontal_line(doc)

    # --- 1. TỔNG QUAN KIẾN TRÚC ---
    doc.add_heading('1. Tổng Quan Kiến Trúc', level=1)

    add_para(doc, 'Mô hình: 2 Sites phân tán giao tiếp qua HTTP REST API (Flask).', size=12)
    add_styled_table(doc,
        ["", "Site A (Trụ sở — port 5000)", "Site B (Web Server — port 5001)"],
        [
            ["Bảng", "Subscribers (1M rows)", "WebLogs (10M rows)"],
            ["Schema", "user_id, full_name, plan, fee", "log_id, user_id, page, status, timestamp"],
            ["Vai trò", "Tạo BF → Gửi → Nhận kết quả → JOIN cuối", "Nhận BF → Lọc WebLogs → Gửi trả"],
        ],
        col_widths=[2, 6, 6]
    )
    doc.add_paragraph()
    add_para(doc, 'Truy vấn: SELECT * FROM Subscribers S JOIN WebLogs W ON S.user_id = W.user_id',
             bold=True, italic=True, size=11)
    add_para(doc, 'Naive Join: ~1,430 MB (gửi hết 10M rows) → BF Semi-Join: ~307 MB → Tiết kiệm ~78.5%',
             size=11, bold=True)

    # --- 2. QUY TRÌNH SEMI-JOIN 5 BƯỚC ---
    doc.add_heading('2. Quy Trình Semi-Join — 5 Bước', level=1)
    add_styled_table(doc,
        ["Bước", "Vị trí", "Hành động", "Chi phí mạng"],
        [
            ["1", "Site A", "Tạo BF từ 1M user_ids (FPR=1%)", "0"],
            ["2", "A → B", "Gửi BF bit-vector qua HTTP", "~1.14 MB"],
            ["3", "Site B", "Lọc WebLogs qua BF, loại ~80% guest", "0"],
            ["4", "B → A", "Gửi filtered logs (~2M + FP rows)", "~306 MB"],
            ["5", "Site A", "Inner Join cuối → loại FP → kết quả 100%", "0"],
        ],
        col_widths=[1.2, 1.8, 7.5, 3]
    )

    # --- 3. BLOOM FILTER ---
    doc.add_heading('3. Bloom Filter — Thiết Kế Cốt Lõi', level=1)
    add_para(doc, 'Cấu trúc: Bit array m bits + k hàm băm (Double Hashing — MurmurHash3, Kirsch & Mitzenmacher 2006).', size=11)
    add_para(doc, 'Công thức tối ưu (Özsu & Valduriez):', bold=True, size=11)
    add_styled_table(doc,
        ["Tham số", "Công thức", "Ý nghĩa"],
        [
            ["m", "m = −(n × ln(p)) / (ln2)²", "Kích thước bit array tối ưu"],
            ["k", "k = (m/n) × ln(2)", "Số hàm băm tối ưu"],
            ["FPR", "FPR = (1 − e^(−kn/m))^k", "Tỷ lệ dương tính giả thực tế"],
        ],
        col_widths=[2, 5.5, 6]
    )
    doc.add_paragraph()
    add_para(doc, 'Đặc điểm: False Negative = 0% (không bỏ sót) · False Positive > 0% (kiểm soát được, bị loại ở Inner Join cuối).', size=11)

    # --- 4. TECHNOLOGY STACK ---
    doc.add_heading('4. Technology Stack', level=1)
    add_styled_table(doc,
        ["Thành phần", "Công nghệ", "Lý do chọn"],
        [
            ["Hash function", "MurmurHash3 (mmh3)", "Nhanh, phân bố đều, non-cryptographic"],
            ["Bit storage", "bitarray", "Tiết kiệm bộ nhớ, O(1) access"],
            ["Data processing", "Pandas", "Mô phỏng database operations"],
            ["Web framework", "Flask", "RESTful API + Web Dashboard"],
            ["Visualization", "Chart.js + Matplotlib", "Biểu đồ trực quan"],
            ["Hashing", "Double Hashing", "Kirsch & Mitzenmacher (2006)"],
        ],
        col_widths=[3, 4, 7]
    )

    # --- 5. FAULT TOLERANCE ---
    doc.add_heading('5. Khả Năng Chịu Lỗi — Kill Node B', level=1)
    add_para(doc, 'Khi Node B bị sập giữa giao dịch JOIN:', bold=True, size=11)
    add_bullet(doc, "Site A gửi BF → Node B không phản hồi → Connection Timeout (HTTP 504).", bold_prefix="Phát hiện: ")
    add_bullet(doc, "Hủy giao dịch (Abort Transaction) → giải phóng bộ nhớ → không cập nhật kết quả.", bold_prefix="Xử lý: ")
    add_bullet(doc, "Dashboard hiển thị pipeline đỏ + banner cảnh báo sập giao dịch phân tán.", bold_prefix="Cảnh báo: ")
    add_bullet(doc, 'Nhấn "Khôi phục Node B" → kiểm tra kết nối → chạy lại bình thường.', bold_prefix="Recovery: ")
    add_para(doc, 'Đây là mô phỏng giao thức Atomic Commitment — đảm bảo tính Atomicity trong hệ phân tán.',
             size=11, italic=True)

    # --- LƯU FILE ---
    filepath = os.path.join(OUTPUT_DIR, "Design_Document_v2.docx")
    doc.save(filepath)
    print(f"  [OK] Da tao: {filepath}")


# ============================================================
# FILE 3: BAO CAO PHAN TICH
# ============================================================

def create_analysis_report():
    doc = Document()
    setup_styles(doc)

    # --- TRANG BIA ---
    add_title_page(doc,
        ["BÁO CÁO PHÂN TÍCH"],
        "Bloom Filter Semi-Join Optimizer — \"Subscribers & Logs\""
    )
    add_para(doc, "Đề tài #15 | Môn: Cơ Sở Dữ Liệu Phân Tán", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # --- 1. GIOI THIEU ---
    doc.add_heading('1. GIỚI THIỆU ĐỀ TÀI', level=1)
    add_horizontal_line(doc)

    doc.add_heading('1.1. Bối cảnh', level=2)
    add_para(doc,
        'Trong hệ thống Cơ Sở Dữ Liệu Phân Tán (CSDL-PT), dữ liệu được phân bố trên nhiều '
        'site (nút mạng) khác nhau. Khi thực hiện phép JOIN giữa các bảng nằm ở các site khác '
        'nhau, hệ thống phải truyền tải dữ liệu qua mạng — đây là chi phí lớn nhất trong xử '
        'lý truy vấn phân tán.', size=12)

    add_para(doc,
        'Theo Özsu & Valduriez (Principles of Distributed Database Systems, Chapter 7):', bold=True, size=12)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('"The cost of transferring data over the network is typically the dominant cost '
                    'in distributed query processing. Therefore, minimizing the amount of data '
                    'transferred is crucial."')
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_heading('1.2. Mô hình dữ liệu', level=2)
    add_styled_table(doc,
        ["Site", "Bảng", "Mô tả", "Quy mô"],
        [
            ["Site A (Trụ sở)", "Subscribers", "Thuê bao trả phí", "1,000,000 rows"],
            ["Site B (Web Server)", "WebLogs", "Nhật ký truy cập", "10,000,000 rows"],
        ],
        col_widths=[3.5, 3, 4, 3.5]
    )
    doc.add_paragraph()
    add_para(doc, 'Truy vấn cần thực thi:', bold=True, size=12)
    add_para(doc, 'SELECT S.full_name, S.plan, W.page, W.timestamp\nFROM Subscribers S JOIN WebLogs W ON S.user_id = W.user_id',
             italic=True, size=11, indent=0.5)

    doc.add_heading('1.3. Vấn đề', level=2)
    add_para(doc,
        'Trong thực tế, phần lớn lượt truy cập website đến từ khách vãng lai (guest), '
        'chỉ khoảng 20-35% là từ thuê bao. Nếu áp dụng JOIN truyền thống (gửi toàn bộ WebLogs '
        'từ Site B → Site A), sẽ lãng phí 65-80% băng thông mạng.', size=12)

    # --- 2. CO SO LY THUYET ---
    doc.add_heading('2. CƠ SỞ LÝ THUYẾT (Theo Özsu & Valduriez)', level=1)
    add_horizontal_line(doc)

    doc.add_heading('2.1. Xử lý truy vấn phân tán', level=2)
    add_para(doc, 'Theo Özsu & Valduriez, Chương 7, quá trình xử lý gồm 4 lớp:', size=12)
    layers = [
        ("Query Decomposition: ", "Phân rã truy vấn thành các phép toán quan hệ."),
        ("Data Localization: ", "Xác định dữ liệu nằm ở site nào."),
        ("Global Optimization: ", "Tối ưu thứ tự thực hiện và chiến lược truyền tải."),
        ("Local Optimization: ", "Tối ưu tại từng site riêng lẻ."),
    ]
    for i, (prefix, desc) in enumerate(layers, 1):
        add_bullet(doc, desc, bold_prefix=f"{i}. {prefix}")
    add_para(doc, 'Đồ án này tập trung vào lớp 3 — Global Optimization, cụ thể là tối ưu chiến '
             'lược truyền tải dữ liệu giữa các site khi thực hiện phép JOIN.', size=12, italic=True)

    doc.add_heading('2.2. Semi-Join (Nửa phép kết)', level=2)
    add_para(doc, 'Theo Mục 7.4.3, Semi-Join giảm lượng dữ liệu truyền tải bằng cách gửi trước '
             'giá trị join attribute (projection) rồi lọc tại remote site.', size=12)

    doc.add_heading('2.3. Bloom Filter — Cải tiến Semi-Join', level=2)
    add_para(doc,
        'Bloom Filter (Bloom, 1970) là cấu trúc dữ liệu xác suất cho phép biểu diễn tập hợp '
        'dưới dạng bit-vector cực nhỏ. Thay vì gửi 10 MB danh sách user_ids, ta gửi Bloom Filter '
        'chỉ ~1.14 MB.', size=12)

    add_para(doc, 'Công thức toán học:', bold=True, size=12)
    add_styled_table(doc,
        ["Tham số", "Công thức", "Ý nghĩa"],
        [
            ["m (kích thước tối ưu)", "m = −(n × ln(p)) / (ln2)²", "Số bits cần cho n phần tử với FPR = p"],
            ["k (số hash tối ưu)", "k = (m/n) × ln(2)", "Cân bằng giữa precision và collision"],
            ["FPR (thực tế)", "FPR = (1 − e^(−kn/m))^k", "Xác suất nhận nhầm phần tử không có"],
        ],
        col_widths=[4, 5, 5.5]
    )
    doc.add_paragraph()

    add_para(doc, 'Đặc điểm quan trọng:', bold=True, size=12)
    add_styled_table(doc,
        ["Đặc điểm", "Giá trị", "Ý nghĩa"],
        [
            ["False Negative Rate", "= 0%", "KHÔNG BAO GIỜ bỏ sót phần tử thực sự có"],
            ["False Positive Rate", "> 0%", "CÓ THỂ nhận nhầm phần tử không có"],
            ["Tính toàn vẹn kết quả", "100%", "FP bị loại ở bước Inner Join cuối cùng"],
        ],
        col_widths=[4.5, 3, 7]
    )

    # --- 3. THIET KE VA CAI DAT ---
    doc.add_heading('3. THIẾT KẾ VÀ CÀI ĐẶT', level=1)
    add_horizontal_line(doc)

    doc.add_heading('3.1. Lựa chọn thiết kế (Design Choices)', level=2)
    add_para(doc, 'Hàm băm: MurmurHash3 + Double Hashing', bold=True, size=12)
    add_para(doc,
        'Dùng kỹ thuật Double Hashing (Kirsch & Mitzenmacher, 2006): '
        'hᵢ(x) = (h₁(x) + i × h₂(x)) mod m, trong đó h₁ và h₂ là hai hàm MurmurHash3 '
        'với seed khác nhau. Giảm từ k lần hash xuống còn 2 lần, hiệu quả tương đương.', size=12)

    add_para(doc, 'Cấu hình tham số tự động', bold=True, size=12)
    add_bullet(doc, "Chỉ cần cung cấp n (số phần tử) và fp_rate → tự tính m và k tối ưu.", bold_prefix="Tự động (recommended): ")
    add_bullet(doc, "Chỉ định trực tiếp m và k.", bold_prefix="Thủ công: ")

    # --- 4. KET QUA ---
    doc.add_heading('4. KẾT QUẢ PHÂN TÍCH', level=1)
    add_horizontal_line(doc)

    # Ghi chú số liệu
    p = doc.add_paragraph()
    run = p.add_run('Lưu ý: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run = p.add_run('Toàn bộ số liệu trong phần này được thu thập từ mô phỏng trên máy đơn (localhost) '
                    'với dữ liệu giả lập. Các giá trị bandwidth, FPR, Leverage là kết quả tính toán '
                    'theo giả định overlap ratio = 20%, kích thước dòng trung bình = 150 bytes. '
                    'Kết quả phục vụ mục đích minh họa lý thuyết và so sánh chiến lược.')
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(100, 100, 100)
    p.paragraph_format.space_after = Pt(8)

    doc.add_heading('4.1. Kết quả mô phỏng (100K / 1M) — Chạy trên máy đơn', level=2)
    add_para(doc, 'Mô phỏng với 100,000 Subscribers + 1,000,000 WebLogs (tỷ lệ 1:10), overlap giả định 20%:', size=12)
    add_styled_table(doc,
        ["Chiến lược", "Dòng gửi", "FP", "Mạng (KB)", "Tiết kiệm"],
        [
            ["Naive Join", "1,000,000", "0", "146,484", "baseline"],
            ["BF Semi-Join (FPR 10%)", "~280,000", "~80,000", "~41,016", "~72%"],
            ["BF Semi-Join (FPR 5%)", "~240,000", "~40,000", "~35,156", "~76%"],
            ["BF Semi-Join (FPR 1%)", "~208,000", "~8,000", "~30,469", "~79%"],
            ["BF Semi-Join (FPR 0.1%)", "~200,800", "~800", "~29,414", "~80%"],
        ],
        col_widths=[4.5, 2.5, 2, 2.5, 2.5]
    )

    doc.add_heading('4.2. METRIC CHÍNH: Bytes Saved vs. Size of Bit-Vector', level=2)
    add_para(doc, 'Đây là metric chính theo yêu cầu đề bài:', bold=True, size=12)
    add_styled_table(doc,
        ["FPR", "BF size (m bits)", "BF size (KB)", "Bytes Saved (KB)", "Leverage"],
        [
            ["10%", "479,253", "58.5", "~105,468", "~1,803x"],
            ["5%", "623,527", "76.1", "~111,328", "~1,463x"],
            ["1%", "958,506", "117.0", "~116,015", "~991x"],
            ["0.1%", "1,437,759", "175.5", "~117,070", "~667x"],
        ],
        col_widths=[2, 3.5, 3, 3.5, 2.5]
    )
    doc.add_paragraph()
    add_para(doc, 'Nhận xét:', bold=True, size=12)
    add_bullet(doc, "Mỗi 1 byte Bloom Filter tiết kiệm được 667 – 1,803 bytes bandwidth.", bold_prefix="Leverage cực cao: ")
    add_bullet(doc, "FPR = 1% cho cân bằng tốt nhất giữa kích thước BF và bandwidth saved.", bold_prefix="Trade-off tối ưu: ")

    doc.add_heading('4.3. Phân tích lý thuyết quy mô đề bài (1M / 10M)', level=2)
    add_para(doc, 'Ghi chú: Bảng dưới đây là kết quả tính toán lý thuyết thuần túy bằng công thức '
             'Bloom Filter (không chạy thực nghiệm ở quy mô này). Giả định: overlap 20%, '
             'kích thước dòng WebLog trung bình 150 bytes.', size=11, italic=True)
    add_styled_table(doc,
        ["FPR", "m (bits)", "BF (MB)", "Total (MB)", "Saved (MB)", "Saved %", "Leverage"],
        [
            ["10%", "4,792,530", "0.57", "401.97", "1,028.03", "71.9%", "1,803x"],
            ["5%", "6,235,224", "0.74", "344.83", "1,085.17", "75.9%", "1,462x"],
            ["1%", "9,585,059", "1.14", "306.54", "1,123.46", "78.6%", "983x"],
            ["0.5%", "11,035,332", "1.31", "297.40", "1,132.60", "79.2%", "862x"],
            ["0.1%", "14,377,588", "1.71", "288.26", "1,141.74", "79.8%", "667x"],
        ],
        col_widths=[1.5, 2.5, 1.8, 2, 2, 1.8, 2]
    )
    doc.add_paragraph()
    add_para(doc, 'Kết luận quy mô 1M/10M:', bold=True, size=12)
    add_bullet(doc, "Naive Join cần truyền ~1,430 MB qua mạng.")
    add_bullet(doc, "BF Semi-Join (FPR=1%) chỉ cần ~307 MB → tiết kiệm ~1,123 MB (78.6%).")
    add_bullet(doc, "Bloom Filter chỉ nặng ~1.14 MB — bằng 0.08% dữ liệu tiết kiệm được.")

    # --- 5. DANH GIA THEO LY THUYET ---
    doc.add_heading('5. ĐÁNH GIÁ THEO LÝ THUYẾT ÖZSU & VALDURIEZ', level=1)
    add_horizontal_line(doc)

    doc.add_heading('5.1. So sánh với các chiến lược trong giáo trình', level=2)
    add_styled_table(doc,
        ["Chiến lược", "Chi phí mạng (1M/10M)", "Ưu điểm", "Nhược điểm"],
        [
            ["Ship Whole (Naive)", "~1,430 MB", "Đơn giản", "Lãng phí cực lớn"],
            ["Semi-Join (gửi IDs)", "~317 MB", "Giảm ~78%", "Vẫn tốn 10 MB gửi IDs"],
            ["BF Semi-Join (đề án)", "~307 MB", "Giảm ~78.6%, BF nhỏ gọn", "FPR > 0 (không ảnh hưởng kết quả)"],
        ],
        col_widths=[4, 3.5, 3.5, 4]
    )

    doc.add_heading('5.2. Chi phí truyền tải (Communication Cost)', level=2)
    add_para(doc, 'Theo Mục 7.3.2, chi phí truyền tải: C_total = C₀ + C₁ × |data|', bold=True, size=12)
    add_styled_table(doc,
        ["Phương pháp", "|data| truyền", "Số lần truyền"],
        [
            ["Naive Join", "10M × 150 bytes = 1,430 MB", "1 chiều (B→A)"],
            ["Semi-Join", "1M × 10 bytes + 2M × 150 bytes ≈ 310 MB", "2 chiều (A→B→A)"],
            ["BF Semi-Join", "1.14 MB + 2M × 150 bytes ≈ 307 MB", "2 chiều, chiều đi rất nhỏ"],
        ],
        col_widths=[3.5, 6, 4]
    )

    doc.add_heading('5.3. Tính đúng đắn (Correctness)', level=2)
    add_para(doc, 'BF Semi-Join đảm bảo tính đúng đắn tuyệt đối:', bold=True, size=12)
    add_bullet(doc, "False Negative = 0% → KHÔNG BAO GIỜ bỏ sót subscriber.")
    add_bullet(doc, "False Positive được loại bỏ → Ở bước Inner Join cuối cùng tại Site A.")
    add_bullet(doc, "Kết quả = 100% chính xác → Tương đương với centralized JOIN.")

    doc.add_heading('5.4. Trade-off Analysis', level=2)
    add_styled_table(doc,
        ["Trade-off", "Phân tích"],
        [
            ["BF size vs FPR", "m lớn → FPR nhỏ → ít FP → tiết kiệm hơn, nhưng BF nặng hơn"],
            ["Precision vs Bandwidth", "FPR=0.1% tiết kiệm 79.8% nhưng BF=1.71 MB; FPR=10% tiết kiệm 71.9% nhưng BF=0.57 MB"],
            ["Computation vs Communication", "Tính hash tại local (rẻ) để giảm data truyền qua mạng (đắt)"],
            ["Optimal point", "FPR = 1% là điểm cân bằng tốt nhất cho bài toán này"],
        ],
        col_widths=[4, 10.5]
    )

    doc.add_heading('5.5. Kịch bản chịu lỗi — Kill Node B (Fault Tolerance)', level=2)
    add_para(doc,
        'Theo Özsu & Valduriez, Chương 12, hệ thống phân tán cần đảm bảo tính nhất quán '
        'ngay cả khi một nút bị sập. Đồ án mô phỏng kịch bản Kill Node B:', size=12)
    add_bullet(doc, "Site A thiết lập Timeout (5s). Quá thời gian hoặc nhận HTTP 504 → xác định Node B đã sập.",
               bold_prefix="1. Phát hiện lỗi: ")
    add_bullet(doc, "Hủy giao dịch (Abort Transaction) → giải phóng bộ nhớ → không cập nhật kết quả thiếu nhất quán.",
               bold_prefix="2. Abort & Rollback: ")
    add_bullet(doc, "Dashboard hiển thị pipeline đỏ + banner cảnh báo sập giao dịch phân tán.",
               bold_prefix="3. Cảnh báo trực quan: ")
    add_bullet(doc, 'Nhấn "Khôi phục Node B" → kiểm tra kết nối → chạy lại bình thường.',
               bold_prefix="4. Recovery: ")
    add_para(doc,
        'Ý nghĩa: Đây là mô phỏng giao thức Atomic Commitment — nếu bất kỳ bước nào thất bại, '
        'toàn bộ giao dịch bị hủy (Abort) chứ không hoàn thành một phần (Partial Commit). '
        'Đảm bảo tính Atomicity trong hệ phân tán.', size=11, italic=True)

    # --- 6. KET LUAN ---
    doc.add_heading('6. KẾT LUẬN', level=1)
    add_horizontal_line(doc)

    doc.add_heading('6.1. Tóm tắt kết quả', level=2)
    add_styled_table(doc,
        ["Metric", "Giá trị"],
        [
            ["Bandwidth tiết kiệm (FPR=1%, 1M/10M)", "78.6% (~1,123 MB saved)"],
            ["Kích thước BF (1M users, FPR=1%)", "1.14 MB (m = 9,585,059 bits)"],
            ["Savings Leverage", "983x (1 byte BF → tiết kiệm 983 bytes)"],
            ["Độ chính xác kết quả", "100% (FP bị loại ở Inner Join)"],
            ["False Negative", "0% (BF không bao giờ bỏ sót subscriber)"],
        ],
        col_widths=[7, 7]
    )

    doc.add_heading('6.2. Đóng góp', level=2)
    add_bullet(doc, "Cài đặt Bloom Filter from scratch — Không dùng thư viện BF có sẵn.")
    add_bullet(doc, "Mô phỏng đầy đủ pipeline Semi-Join — 5 bước, đo lường chi tiết bandwidth.")
    add_bullet(doc, "Phân tích cả lý thuyết lẫn thực nghiệm — So sánh FPR lý thuyết vs thực tế.")
    add_bullet(doc, "Web Dashboard tương tác — Cho phép tuỳ chỉnh tham số và xem kết quả trực quan.")
    add_bullet(doc, "Liên kết chặt chẽ với giáo trình Özsu & Valduriez — Đúng framework Distributed Query Processing.")

    doc.add_heading('6.3. Hạn chế và hướng phát triển', level=2)
    add_styled_table(doc,
        ["Hạn chế", "Hướng phát triển"],
        [
            ["Demo 100K/1M (thay vì full 1M/10M)", "Dùng streaming/chunked processing"],
            ["Mô phỏng trên 1 máy", "Triển khai distributed thực sự (Docker, gRPC)"],
            ["Bloom Filter cơ bản", "Counting Bloom Filter, Cuckoo Filter"],
            ["Chỉ hỗ trợ equi-join", "Mở rộng cho range join, multi-attribute join"],
        ],
        col_widths=[7, 7]
    )

    # --- TAI LIEU THAM KHAO ---
    doc.add_heading('TÀI LIỆU THAM KHẢO', level=1)
    add_horizontal_line(doc)
    refs = [
        'Özsu, M.T. & Valduriez, P. (2020). Principles of Distributed Database Systems, 4th Edition. Springer. — Chương 7: Distributed Query Processing.',
        'Bloom, B.H. (1970). Space/time trade-offs in hash coding with allowable errors. Communications of the ACM, 13(7), 422-426.',
        'Kirsch, A. & Mitzenmacher, M. (2006). Less Hashing, Same Performance: Building a Better Bloom Filter. Proceedings of the 14th Annual European Symposium on Algorithms (ESA).',
        'Mackert, L.F. & Lohman, G.M. (1986). R* optimizer validation and performance evaluation for distributed queries. Proceedings of the 12th International Conference on Very Large Data Bases.',
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run_n = p.add_run(f"[{i}] ")
        run_n.bold = True
        run_n.font.size = Pt(11)
        run_n.font.name = 'Times New Roman'
        run_t = p.add_run(ref)
        run_t.font.size = Pt(11)
        run_t.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(4)

    # --- LUU FILE ---
    filepath = os.path.join(OUTPUT_DIR, "BaoCao_PhanTich_v2.docx")
    doc.save(filepath)
    print(f"  [OK] Da tao: {filepath}")


# ============================================================
# MAIN
# ============================================================

if __name__ == '__main__':
    print("\n" + "=" * 60)
    print("  TAO FILE .DOCX CHUAN CHO DO AN CSDL PHAN TAN")
    print("=" * 60 + "\n")

    create_project_proposal()
    create_design_document()
    create_analysis_report()

    print("\n" + "=" * 60)
    print("  HOAN THANH! Da tao 3 file .docx:")
    print("    1. Project_Proposal.docx")
    print("    2. Design_Document.docx")
    print("    3. BaoCao_PhanTich.docx")
    print("=" * 60 + "\n")
